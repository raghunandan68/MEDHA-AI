import os
import re
import csv
import uuid
import base64
import io
import logging
from pathlib import Path

from app.config import settings

logger = logging.getLogger(__name__)

UPLOAD_DIR = Path(__file__).resolve().parent.parent.parent / "uploads"

SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}
SUPPORTED_DOC_EXTENSIONS = {".docx", ".xlsx", ".csv", ".txt"}
ALL_SUPPORTED_EXTENSIONS = SUPPORTED_IMAGE_EXTENSIONS | SUPPORTED_DOC_EXTENSIONS | {".pdf"}


def ensure_upload_dir():
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


def save_upload(file_bytes: bytes, original_filename: str) -> tuple[str, str]:
    ensure_upload_dir()
    ext = Path(original_filename).suffix or ".pdf"
    stored_name = f"{uuid.uuid4()}{ext}"
    file_path = UPLOAD_DIR / stored_name
    with open(file_path, "wb") as f:
        f.write(file_bytes)
    return stored_name, str(file_path)


def _get_tesseract_langs() -> str:
    return "eng"


def _preprocess_image_for_ocr(image_bytes: bytes) -> bytes:
    try:
        from PIL import Image, ImageFilter, ImageEnhance

        img = Image.open(io.BytesIO(image_bytes))

        if img.mode != "L":
            img = img.convert("L")

        img = img.filter(ImageFilter.SHARPEN)

        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.8)

        enhancer = ImageEnhance.Brightness(img)
        img = enhancer.enhance(1.3)

        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(2.0)

        if img.width < 1500:
            scale = 1500 / img.width
            new_width = int(img.width * scale)
            new_height = int(img.height * scale)
            img = img.resize((new_width, new_height), Image.LANCZOS)

        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception as e:
        logger.warning(f"Image preprocessing failed: {e}")
        return image_bytes


def extract_text_from_image_with_tesseract(image_bytes: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image

        processed = _preprocess_image_for_ocr(image_bytes)
        img = Image.open(io.BytesIO(processed))

        custom_config = r"--oem 3 --psm 6"
        text = pytesseract.image_to_string(img, lang=_get_tesseract_langs(), config=custom_config)
        return text.strip() if text else ""
    except ImportError:
        logger.warning("pytesseract not installed. Install with: pip install pytesseract")
        return ""
    except Exception as e:
        logger.warning(f"Tesseract OCR failed: {e}")
        return ""


VISION_CAPABLE_MODELS = {
    "gpt-4o", "gpt-4o-mini", "gpt-4-vision-preview",
    "claude-3-opus", "claude-3-sonnet", "claude-3-haiku",
    "meta-llama/llama-4-scout-17b-16e-instruct",
    "qwen/qwen3.6-27b",
}


def extract_text_from_image_llm(image_bytes: bytes) -> str:
    groq_key = settings.groq_api_key
    openai_key = settings.openai_api_key

    if groq_key:
        model = settings.ai_model
        if model not in VISION_CAPABLE_MODELS:
            if "groq" in groq_key.lower() or groq_key.startswith("gsk_"):
                model = "meta-llama/llama-4-scout-17b-16e-instruct"
            else:
                return ""
        base_url = "https://api.groq.com/openai/v1"
        api_key = groq_key
    elif openai_key and openai_key != "sk-your-openai-api-key":
        base_url = None
        api_key = openai_key
        model = "gpt-4o-mini"
    else:
        return ""

    try:
        from openai import OpenAI
        client = OpenAI(api_key=api_key, base_url=base_url, timeout=60.0)

        processed = _preprocess_image_for_ocr(image_bytes)
        base64_image = base64.b64encode(processed).decode("utf-8")

        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "You are an expert OCR assistant. Extract ALL visible text from this image with maximum accuracy. "
                                "Preserve the original structure, formatting, paragraphs, bullet points, headers, and tables as closely as possible. "
                                "Return ONLY the extracted text content. Do not add any commentary, descriptions, or explanations. "
                                "If the image contains a table, reproduce it in a structured format. "
                                "If the image contains mathematical expressions, write them in plain text form."
                            )
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/png;base64,{base64_image}"
                            }
                        }
                    ]
                }
            ],
            max_tokens=4000
        )
        return response.choices[0].message.content or ""
    except Exception as e:
        logger.warning(f"LLM vision OCR failed: {e}")
        return ""


def extract_text_from_image(image_bytes: bytes) -> str:
    tesseract_text = extract_text_from_image_with_tesseract(image_bytes)

    llm_text = extract_text_from_image_llm(image_bytes)

    if llm_text and len(llm_text.strip()) > 20:
        return llm_text

    if tesseract_text and len(tesseract_text.strip()) > 20:
        return tesseract_text

    if llm_text and llm_text.strip():
        return llm_text
    if tesseract_text and tesseract_text.strip():
        return tesseract_text
    return ""


def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        return "Word text extraction requires python-docx. Install with: pip install python-docx"

    try:
        doc = Document(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        if not paragraphs:
            return "No text could be extracted from this Word document."

        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error extracting text from Word document: {e}")
        return f"Error extracting text from Word document: {e}"


def extract_text_from_xlsx(file_path: str) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "Excel text extraction requires openpyxl. Install with: pip install openpyxl"

    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)
        all_text = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_lines = []
            for row in ws.iter_rows(values_only=True):
                row_vals = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if row_vals:
                    sheet_lines.append(" | ".join(row_vals))
            if sheet_lines:
                all_text.append(f"Sheet: {sheet_name}\n" + "\n".join(sheet_lines))

        wb.close()

        if not all_text:
            return "No text could be extracted from this Excel file."

        return "\n\n".join(all_text)
    except Exception as e:
        logger.error(f"Error extracting text from Excel file: {e}")
        return f"Error extracting text from Excel file: {e}"


def extract_text_from_csv(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        lines = content.split("\n")
        parsed_rows = []
        try:
            reader = csv.reader(io.StringIO(content))
            for row in reader:
                if any(cell.strip() for cell in row):
                    parsed_rows.append(" | ".join(cell.strip() for cell in row if cell.strip()))
        except csv.Error:
            parsed_rows = [line.strip() for line in lines if line.strip()]

        if not parsed_rows:
            return "No text could be extracted from this CSV file."

        return "\n".join(parsed_rows)
    except Exception as e:
        logger.error(f"Error extracting text from CSV file: {e}")
        return f"Error extracting text from CSV file: {e}"


def extract_images_from_pdf(file_path: str) -> list[bytes]:
    try:
        import fitz
    except ImportError:
        return []

    images = []
    try:
        doc = fitz.open(file_path)
        for page_num in range(len(doc)):
            page = doc[page_num]
            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                if image_bytes and len(image_bytes) > 500:
                    images.append(image_bytes)
        doc.close()
    except Exception as e:
        logger.warning(f"Error extracting images from PDF: {e}")
    return images


def render_pdf_pages_as_images(file_path: str, dpi: int = 200) -> list[bytes]:
    try:
        import fitz
    except ImportError:
        return []

    page_images = []
    try:
        doc = fitz.open(file_path)
        zoom = dpi / 72
        matrix = fitz.Matrix(zoom, zoom)

        for page_num in range(len(doc)):
            page = doc[page_num]
            pix = page.get_pixmap(matrix=matrix, alpha=False)

            if pix.width < 100 or pix.height < 100:
                continue

            img_bytes = pix.tobytes("png")
            if img_bytes and len(img_bytes) > 1000:
                page_images.append(img_bytes)

        doc.close()
    except Exception as e:
        logger.warning(f"Error rendering PDF pages: {e}")
    return page_images


def _has_embedded_text(file_path: str) -> bool:
    try:
        import fitz
        doc = fitz.open(file_path)
        for page in doc:
            text = page.get_text().strip()
            if len(text) > 50:
                doc.close()
                return True
        doc.close()
        return False
    except Exception:
        return False


def extract_text_from_pdf(file_path: str) -> str:
    try:
        import fitz
    except ImportError:
        return "PDF text extraction requires PyMuPDF. Install with: pip install PyMuPDF"

    try:
        doc = fitz.open(file_path)
        page_texts = []
        for page in doc:
            text = page.get_text()
            if text and text.strip():
                page_texts.append(text.strip())
        doc.close()

        embedded_text = "\n\n".join(page_texts)

        has_text = _has_embedded_text(file_path)

        image_texts = []
        images = extract_images_from_pdf(file_path)
        for img_bytes in images[:5]:
            img_text = extract_text_from_image(img_bytes)
            if img_text and img_text.strip():
                image_texts.append(img_text.strip())

        if not has_text and not image_texts:
            logger.info("No embedded text found. Rendering pages for OCR...")
            page_images = render_pdf_pages_as_images(file_path, dpi=200)
            for page_img in page_images[:5]:
                page_ocr_text = extract_text_from_image(page_img)
                if page_ocr_text and page_ocr_text.strip():
                    image_texts.append(page_ocr_text.strip())

        if image_texts:
            combined_image_text = "\n\n".join(image_texts)
            if embedded_text:
                text = embedded_text + "\n\n" + combined_image_text
            else:
                text = combined_image_text
        else:
            text = embedded_text

        if not text or not text.strip():
            return "No text could be extracted from this PDF. The document may be a scanned image without OCR capability."

        return clean_extracted_text(text)
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return f"Error extracting text: {e}"


def extract_text_from_image_file(file_path: str) -> str:
    try:
        with open(file_path, "rb") as f:
            image_bytes = f.read()

        text = extract_text_from_image(image_bytes)
        if text and text.strip():
            return text

        tesseract_text = extract_text_from_image_with_tesseract(image_bytes)
        if tesseract_text and tesseract_text.strip():
            return tesseract_text

        return "No text could be extracted from this image."
    except Exception as e:
        logger.error(f"Error extracting text from image: {e}")
        return f"Error extracting text from image: {e}"


def save_text(text_content: str, filename: str = "pasted.txt") -> tuple[str, str]:
    ensure_upload_dir()
    stored_name = f"{uuid.uuid4()}.txt"
    file_path = UPLOAD_DIR / stored_name
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(text_content)
    return stored_name, str(file_path)


def clean_extracted_text(text: str) -> str:
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned.append(line)
            continue
        lower = stripped.lower()
        if re.search(r"^\s*page\s*\d+\s*(?:of|/)\s*\d+\s*$", lower, re.IGNORECASE):
            continue
        if re.search(r"^\s*\d+\s*/\s*\d+\s*$", stripped):
            continue
        if re.search(r"^\s*\d+\s*$", stripped) and len(stripped) <= 4:
            continue
        if re.search(r"^\s*page\s*\d+\s*$", lower, re.IGNORECASE):
            continue
        cleaned.append(line)

    result = "\n".join(cleaned)
    result = re.sub(r"\n{3,}", "\n\n", result)
    return result.strip()


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".txt":
        try:
            return path.read_text(encoding="utf-8").strip() or "No text content."
        except Exception as e:
            return f"Error reading text file: {e}"

    if ext in SUPPORTED_IMAGE_EXTENSIONS:
        return extract_text_from_image_file(file_path)

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)

    if ext == ".docx":
        return extract_text_from_docx(file_path)

    if ext == ".xlsx":
        return extract_text_from_xlsx(file_path)

    if ext == ".csv":
        return extract_text_from_csv(file_path)

    return f"Unsupported file format: {ext}"


def cleanup_file(file_path: str):
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception:
        pass
